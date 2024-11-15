import os
import streamlit as st
import concurrent.futures
from langchain.docstore.document import Document
from langchain.document_loaders import TextLoader
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from functools import lru_cache

DOCS_DIR = os.path.abspath("./uploaded_docs")
if not os.path.exists(DOCS_DIR):
    os.makedirs(DOCS_DIR)

def load_txt_file(file_path):
    loader = TextLoader(file_path)
    return loader.load()

def load_pdf_file(file_path):
    pages = []
    with open(file_path, "rb") as f:
        pdf_reader = PdfReader(f)
        for page in pdf_reader.pages:
            content = page.extract_text()
            if content:
                pages.append(Document(page_content=content))
    return pages

def load_docx_file(file_path):
    doc = DocxDocument(file_path)
    content = "\n".join([para.text for para in doc.paragraphs if para.text])
    return [Document(page_content=content)]

# Caching the document loading
@lru_cache(maxsize=1)
def load_documents_from_directory():
    raw_documents = []
    if os.path.exists(DOCS_DIR):
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = []
            for file in os.listdir(DOCS_DIR):
                file_path = os.path.join(DOCS_DIR, file)
                if file.endswith('.txt'):
                    futures.append(executor.submit(load_txt_file, file_path))
                elif file.endswith('.pdf'):
                    futures.append(executor.submit(load_pdf_file, file_path))
                elif file.endswith('.docx'):
                    futures.append(executor.submit(load_docx_file, file_path))

            for future in concurrent.futures.as_completed(futures):
                raw_documents.extend(future.result())
    return raw_documents